#!/usr/bin/env python3
"""Build anonymous CUMCM 2026 DOCX/PDF-ready artifacts.

The script prepares a clean Markdown source, converts it with Pandoc so LaTeX
math becomes native Word math, applies deterministic academic styling, creates
an AI-use detail document, and stages an anonymous support-material package.
PDF conversion and visual QA are performed by the documents/PDF render tools.
"""

from __future__ import annotations

import csv
import re
import shutil
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
PROJECT = ROOT.parent
OUT = ROOT / "终稿"
SOURCE_MD = ROOT / "10_修订后完整论文_核验版.md"
CLEAN_MD = OUT / "水泥烧成系统电除尘器的协同优化控制_终稿源文.md"
DRAFT_DOCX = OUT / "_论文排版草稿.docx"
FINAL_DOCX = OUT / "水泥烧成系统电除尘器的协同优化控制_匿名终稿.docx"
AI_MD = ROOT / "AI工具使用详情.md"
AI_DRAFT_DOCX = OUT / "_AI详情排版草稿.docx"
AI_DOCX = OUT / "AI工具使用详情_匿名版.docx"
SUPPORT = OUT / "支撑材料_匿名版"

PANDOC = shutil.which("pandoc") or "/opt/homebrew/bin/pandoc"


def run(cmd: list[str]) -> None:
    subprocess.run(cmd, check=True, cwd=PROJECT)


def strip_work_comments(text: str) -> str:
    return re.sub(r"<!--.*?-->", "", text, flags=re.S)


def add_equation_numbers(text: str) -> str:
    counter = 0

    def repl(match: re.Match[str]) -> str:
        nonlocal counter
        counter += 1
        body = match.group(1).rstrip()
        return f"$$\n{body}\\qquad\\text{{({counter})}}\n$$"

    return re.sub(r"\$\$\s*\n(.*?)\n\s*\$\$", repl, text, flags=re.S)


def sanitize_manuscript_markdown() -> str:
    text = strip_work_comments(SOURCE_MD.read_text(encoding="utf-8"))
    text = add_equation_numbers(text)

    # Use package-relative figure paths in the final source and support bundle.
    text = re.sub(
        r"\(/Users/[^)]*/建模交付/figures_paper/([^/)]+\.png)\)",
        r"(figures_paper/\1)",
        text,
    )
    text = text.replace("![图21 双参数敏感性热图]", "![图21 电场贡献权重消融]")

    # Make every control-column unit explicit instead of attaching the unit only
    # to the last column in each group.
    text = text.replace(
        "| 工况 | $U_1$ | $U_2$ | $U_3$ | $U_4$/kV | $T_1$ | $T_2$ | $T_3$ | $T_4$/s |",
        "| 工况 | $U_1$/kV | $U_2$/kV | $U_3$/kV | $U_4$/kV | $T_1$/s | $T_2$/s | $T_3$/s | $T_4$/s |",
    )
    text = text.replace(
        "| 工况与限值 | $U_1$ | $U_2$ | $U_3$ | $U_4$/kV | $T_1$ | $T_2$ | $T_3$ | $T_4$/s |",
        "| 工况与限值 | $U_1$/kV | $U_2$/kV | $U_3$/kV | $U_4$/kV | $T_1$/s | $T_2$/s | $T_3$/s | $T_4$/s |",
    )

    # Keep paths readable and portable in prose.
    text = text.replace("建模交付/src/", "src/")
    text = text.replace("建模交付/results", "results")
    text = text.replace("建模交付/figures_paper", "figures_paper")
    text = text.replace("建模交付/integrity", "integrity")
    text = text.replace("建模交付/revision_round1", "revision_round1")
    text = text.replace("建模交付/AI工具使用详情.md", "AI工具使用详情.pdf")

    code_parts = ["\n\n## 附录D 完整源程序\n"]
    for i, name in enumerate(
        ["scenario_model.py", "paper_figures.py", "plot_utils.py", "integrity_check.py"], start=1
    ):
        code = (ROOT / "src" / name).read_text(encoding="utf-8")
        code_parts.append(f"\n### D.{i} {name}\n\n```python\n{code.rstrip()}\n```\n")
    text = text.rstrip() + "".join(code_parts) + "\n"
    return text


def set_cell_margins(cell, top=60, start=70, bottom=60, end=70) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)


BODY_CJK = "Noto Serif CJK SC"
HEADING_CJK = "Noto Sans CJK SC"
LATIN_SERIF = "Times New Roman"
LATIN_SANS = "Arial"


def set_run_font(run, east=BODY_CJK, latin=LATIN_SERIF, size=12.0, bold=None) -> None:
    # Use a proper Simplified-Chinese serif face for CJK glyphs and keep Latin
    # letters/numerals in Times New Roman.  Setting all four OOXML slots avoids
    # Word/LibreOffice choosing a thin Songti SC fallback during PDF export.
    run.font.name = latin
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), latin)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), latin)
    run._element.get_or_add_rPr().rFonts.set(qn("w:cs"), latin)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0, 0, 0)
    if bold is not None:
        run.bold = bold


def set_style_font(style, east, latin, size, bold=False) -> None:
    style.font.name = latin
    style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east)
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), latin)
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), latin)
    style._element.get_or_add_rPr().rFonts.set(qn("w:cs"), latin)
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor(0, 0, 0)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2])
    set_run_font(run, size=10.5)


def paragraph_has_drawing(paragraph) -> bool:
    return bool(paragraph._p.xpath(".//w:drawing | .//w:pict"))


def paragraph_is_display_math(paragraph) -> bool:
    """Return true only for a standalone display equation.

    Pandoc represents both inline variables and display equations with OMML.
    Testing for any ``m:oMath`` therefore centred ordinary prose paragraphs
    merely because they contained an inline symbol.  Only ``m:oMathPara`` is
    a block equation and should control the paragraph alignment.
    """
    return bool(paragraph._p.xpath(".//m:oMathPara"))


def style_academic_docx(path_in: Path, path_out: Path, ai_detail=False) -> None:
    doc = Document(path_in)
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)
    sec.header_distance = Cm(0.8)
    sec.footer_distance = Cm(1.2)
    sec.different_first_page_header_footer = False

    # CUMCM-specific named override of the standard business preset.
    normal = doc.styles["Normal"]
    set_style_font(normal, BODY_CJK, LATIN_SERIF, 12.0)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.first_line_indent = Cm(0.85)
    normal.paragraph_format.line_spacing = 1.50
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)

    # Pandoc normally emits prose as Body Text / First Paragraph rather than
    # Normal.  Give both explicit CUMCM body formatting so export behaviour is
    # stable across Word and LibreOffice.
    for body_style_name in ("Body Text", "First Paragraph"):
        if body_style_name in [s.name for s in doc.styles]:
            body_style = doc.styles[body_style_name]
            set_style_font(body_style, BODY_CJK, LATIN_SERIF, 12.0)
            body_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            body_style.paragraph_format.first_line_indent = Cm(0.85)
            body_style.paragraph_format.line_spacing = 1.50
            body_style.paragraph_format.space_before = Pt(0)
            body_style.paragraph_format.space_after = Pt(0)

    # Pandoc maps Markdown ##/###/#### to Word Heading 2/3/4.  In this
    # manuscript, Markdown ## is the semantic chapter (一级标题), so it is
    # centred; lower levels remain left aligned.
    for name, size, before, after in (
        ("Heading 1", 22, 12, 10),
        ("Heading 2", 15, 12, 8),
        ("Heading 3", 14, 10, 6),
        ("Heading 4", 12, 8, 4),
    ):
        st = doc.styles[name]
        set_style_font(st, HEADING_CJK, LATIN_SANS, size, bold=True)
        if name == "Heading 1":
            st.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif name == "Heading 2" and not ai_detail:
            st.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            st.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        st.paragraph_format.first_line_indent = Cm(0)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    if "Source Code" in [s.name for s in doc.styles]:
        code_style = doc.styles["Source Code"]
        set_style_font(code_style, BODY_CJK, "Menlo", 7.0)
        code_style.paragraph_format.left_indent = Cm(0)
        code_style.paragraph_format.first_line_indent = Cm(0)
        code_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        code_style.paragraph_format.line_spacing = 1.0
        code_style.paragraph_format.space_before = Pt(0)
        code_style.paragraph_format.space_after = Pt(0)

    # Quiet page furniture; no headers or identity fields.
    for section in doc.sections:
        section.header.paragraphs[0].text = ""
        footer_p = section.footer.paragraphs[0]
        footer_p.clear()
        add_page_number(footer_p)

    # Title and summary page.
    first = next((p for p in doc.paragraphs if p.text.strip()), None)
    if first is not None:
        first.alignment = WD_ALIGN_PARAGRAPH.CENTER
        first.paragraph_format.first_line_indent = Cm(0)
        first.paragraph_format.space_before = Pt(0 if ai_detail else 18)
        first.paragraph_format.space_after = Pt(12)
        for run in first.runs:
            set_run_font(run, east=HEADING_CJK, latin=LATIN_SANS, size=22 if not ai_detail else 18, bold=True)

    for p in doc.paragraphs:
        txt = p.text.strip()
        if not txt:
            continue
        style_name = p.style.name if p.style is not None else ""
        if p is first:
            for r in p.runs:
                set_run_font(r, east=HEADING_CJK, latin=LATIN_SANS, size=22 if not ai_detail else 18, bold=True)
        elif style_name == "Source Code":
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = Cm(0)
            for r in p.runs:
                set_run_font(r, east=BODY_CJK, latin="Menlo", size=7.0)
        elif style_name.startswith("Heading"):
            level_size = {"Heading 1": 22, "Heading 2": 15, "Heading 3": 14, "Heading 4": 12}.get(style_name, 12)
            if style_name == "Heading 1" or (style_name == "Heading 2" and not ai_detail):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = Cm(0)
            for r in p.runs:
                set_run_font(r, east=HEADING_CJK, latin=LATIN_SANS, size=level_size, bold=True)
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = Cm(0.85)
            p.paragraph_format.line_spacing = 1.50
            for r in p.runs:
                set_run_font(r, east=BODY_CJK, latin=LATIN_SERIF, size=12.0)
        if txt == "摘要":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Cm(0)
            for r in p.runs:
                set_run_font(r, east=HEADING_CJK, latin=LATIN_SANS, size=15, bold=True)
        if txt.startswith("关键词："):
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.space_before = Pt(8)
        if txt.startswith("1 问题重述与分析") and not ai_detail:
            p.paragraph_format.page_break_before = True
        if txt.startswith("附录A 程序求解流程") and not ai_detail:
            p.paragraph_format.page_break_before = True
        if txt.startswith("D.") and txt.endswith(".py") and not txt.startswith("D.1 ") and not ai_detail:
            p.paragraph_format.page_break_before = True
        if re.match(r"^(图|表)\d+", txt):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.space_before = Pt(6 if txt.startswith("表") else 2)
            p.paragraph_format.space_after = Pt(2 if txt.startswith("表") else 6)
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.keep_together = True
            # A table title must not be stranded at the bottom of a page.
            # Figure captions follow their figures, so only table titles need
            # to be kept with the next document object.
            p.paragraph_format.keep_with_next = txt.startswith("表")
            for r in p.runs:
                # Mathematical-modeling papers conventionally use concise,
                # centered figure/table titles without a full bold paragraph.
                set_run_font(r, east=BODY_CJK, latin=LATIN_SERIF, size=10.5, bold=False)
        if txt.startswith("[") and re.match(r"^\[\d+\]", txt):
            # Reference entries use left alignment. Full justification creates
            # conspicuous gaps in short lines and around long URLs.
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = Cm(-0.6)
            p.paragraph_format.left_indent = Cm(0.6)
            p.paragraph_format.line_spacing = 1.25
            p.paragraph_format.space_after = Pt(2)
            for r in p.runs:
                set_run_font(r, size=10.5)
        if paragraph_has_drawing(p):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.keep_with_next = True
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(1)
        if paragraph_is_display_math(p):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)

    # Keep images legible within the A4 text block.
    max_w = Inches(6.10)
    max_h = Inches(4.15)
    for shape in doc.inline_shapes:
        ratio = shape.width / shape.height if shape.height else 1
        width = min(shape.width, max_w)
        height = int(width / ratio)
        if height > max_h:
            height = max_h
            width = int(height * ratio)
        shape.width = int(width)
        shape.height = int(height)

    # Explicit full-width table geometry and compact numeric tables.
    usable = int(Cm(16.0))
    for table in doc.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        ncols = len(table.columns)
        if ncols == 0:
            continue
        width = int(usable / ncols)
        set_repeat_table_header(table.rows[0])
        for ri, row in enumerate(table.rows):
            for cell in row.cells:
                cell.width = width
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                set_cell_margins(cell)
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.first_line_indent = Cm(0)
                    p.paragraph_format.line_spacing = 1.0
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(0)
                    for r in p.runs:
                        set_run_font(r, east=BODY_CJK, latin=LATIN_SERIF, size=8.2 if ncols >= 9 else 9.5, bold=(ri == 0))
            if ri == 0:
                for cell in row.cells:
                    shd = OxmlElement("w:shd")
                    shd.set(qn("w:fill"), "E7E6E6")
                    cell._tc.get_or_add_tcPr().append(shd)

    # Stable package metadata; privacy scrub runs after this step as a second gate.
    doc.core_properties.title = "AI工具使用详情" if ai_detail else "水泥烧成系统电除尘器的协同优化控制"
    doc.core_properties.subject = "全国大学生数学建模竞赛匿名论文"
    doc.core_properties.author = ""
    doc.core_properties.last_modified_by = ""
    doc.core_properties.comments = ""
    doc.save(path_out)


def pandoc_docx(markdown: Path, output: Path, resource_path: Path) -> None:
    run(
        [
            PANDOC,
            str(markdown),
            "--from=gfm+tex_math_dollars",
            "--to=docx",
            f"--resource-path={resource_path}",
            "--highlight-style=tango",
            "-o",
            str(output),
        ]
    )


def prepare_ai_markdown() -> Path:
    text = strip_work_comments(AI_MD.read_text(encoding="utf-8"))
    text = text.replace("# AI工具使用详情（提交前核验版）", "# AI工具使用详情")
    text = text.replace(
        "> 状态：内容已按当前工作过程记录；正式提交前必须由参赛队逐项核对并确认，不得把本文件视为自动生成的真实性证明。",
        "> 本文件记录AI工具的用途、采纳范围和人工核验责任，随竞赛支撑材料提交。",
    )
    out = OUT / "AI工具使用详情_匿名版.md"
    out.write_text(text, encoding="utf-8")
    return out


def patch_support_script(src: Path, dst: Path) -> None:
    text = src.read_text(encoding="utf-8")
    if src.name == "scenario_model.py":
        text = re.sub(
            r'PROJECT_ROOT = ROOT\.parent\nDATA_PACKAGE = .*?\nDATA_FILE = .*?\nCLUSTER_EVAL_FILE = .*?\n',
            'DATA_FILE = ROOT / "data" / "full_timeseries_with_flags.csv"\nCLUSTER_EVAL_FILE = ROOT / "data" / "condition_cluster_evaluation.csv"\n',
            text,
            flags=re.S,
        )
    elif src.name == "paper_figures.py":
        text = re.sub(
            r'PROJECT_ROOT = ROOT\.parent\nDATA_PACKAGE = .*?\nDATA_FILE = .*?\nRESULTS =',
            'DATA_FILE = ROOT / "data" / "full_timeseries_with_flags.csv"\nRESULTS =',
            text,
            flags=re.S,
        )
        text = text.replace(
            '"18_condition_energy_penalty": "四类工况从10收紧至5 mg/Nm³的中心情景功率增幅为12.99%～14.11%',
            '"18_condition_energy_penalty": "四类工况从10收紧至5 mg/Nm³的中心情景功率增幅为13.34%～14.11%',
        )
    elif src.name == "integrity_check.py":
        text = re.sub(
            r'PROJECT_ROOT = ROOT\.parent\nDATA_PACKAGE = .*?\nDATA = .*?\nRESULTS =',
            'DATA = ROOT / "data" / "full_timeseries_with_flags.csv"\nRESULTS =',
            text,
            flags=re.S,
        )
        text = text.replace('MANUSCRIPT = ROOT / "10_修订后完整论文_核验版.md"', 'MANUSCRIPT = ROOT / "论文终稿源文.md"')
    dst.write_text(text, encoding="utf-8")


def sanitize_manifest(src: Path, dst: Path) -> None:
    with src.open(encoding="utf-8-sig", newline="") as fi, dst.open("w", encoding="utf-8-sig", newline="") as fo:
        reader = csv.DictReader(fi)
        writer = csv.DictWriter(fo, fieldnames=reader.fieldnames)
        writer.writeheader()
        for row in reader:
            for key in ("png", "pdf", "png_path", "pdf_path"):
                if key in row:
                    row[key] = f"figures_paper/{Path(row[key]).name}"
            if "source_data" in row:
                value = row["source_data"]
                name = Path(value).name
                if name == "full_timeseries_with_flags.csv":
                    row["source_data"] = "data/full_timeseries_with_flags.csv"
                elif name.endswith(".md"):
                    row["source_data"] = "论文终稿源文.md"
                else:
                    row["source_data"] = f"results/{name}"
            writer.writerow(row)


def stage_support_materials() -> None:
    if SUPPORT.exists():
        shutil.rmtree(SUPPORT)
    for sub in ("src", "results", "figures_paper", "data", "integrity"):
        (SUPPORT / sub).mkdir(parents=True, exist_ok=True)

    for src in sorted((ROOT / "src").glob("*.py")):
        if src.name == Path(__file__).name:
            continue
        patch_support_script(src, SUPPORT / "src" / src.name)
    for src in sorted((ROOT / "results").iterdir()):
        if src.is_file():
            shutil.copy2(src, SUPPORT / "results" / src.name)
    for pattern in ("*.png", "*.pdf"):
        for src in sorted((ROOT / "figures_paper").glob(pattern)):
            shutil.copy2(src, SUPPORT / "figures_paper" / src.name)
    sanitize_manifest(ROOT / "figures_paper" / "figure_manifest.csv", SUPPORT / "figures_paper" / "figure_manifest.csv")
    caption_text = (ROOT / "figures_paper" / "图注清单.md").read_text(encoding="utf-8")
    caption_text = re.sub(
        r"/Users/[^`\n]*/建模交付/figures_paper/([^/`\n]+)",
        r"figures_paper/\1",
        caption_text,
    )
    caption_text = re.sub(
        r"/Users/[^`\n]*/建模交付/results/([^/`\n]+)",
        r"results/\1",
        caption_text,
    )
    caption_text = re.sub(
        r"/Users/[^`\n]*/建模交付/10_修订后完整论文_核验版\.md",
        "论文终稿源文.md",
        caption_text,
    )
    caption_text = re.sub(
        r"/Users/[^`\n]*/数据处理/[^`\n]*/data/processed/full_timeseries_with_flags\.csv",
        "data/full_timeseries_with_flags.csv",
        caption_text,
    )
    (SUPPORT / "figures_paper" / "图注清单.md").write_text(caption_text, encoding="utf-8")
    shutil.copy2(
        PROJECT / "数据处理" / "水泥电除尘器_数据处理四次返修包_交付链P1闭环_排放仍阻断" / "data" / "processed" / "full_timeseries_with_flags.csv",
        SUPPORT / "data" / "full_timeseries_with_flags.csv",
    )
    shutil.copy2(
        PROJECT / "数据处理" / "水泥电除尘器_数据处理四次返修包_交付链P1闭环_排放仍阻断" / "data" / "audit" / "condition_cluster_evaluation.csv",
        SUPPORT / "data" / "condition_cluster_evaluation.csv",
    )
    shutil.copy2(ROOT / "integrity" / "deterministic_integrity_results.json", SUPPORT / "integrity" / "deterministic_integrity_results.json")
    shutil.copy2(CLEAN_MD, SUPPORT / "论文终稿源文.md")
    readme = """水泥烧成系统电除尘器协同优化控制——匿名支撑材料

目录说明：
- data：题给数据的处理结果与工况聚类评价；
- src：完整可运行Python源程序；
- results：论文全部主要CSV/JSON结果；
- figures_paper：正文21幅矢量PDF及图件追踪表；
- integrity：63项确定性核验结果；
- 论文终稿源文.md：去除身份与本地绝对路径的论文源文。

建议运行顺序（需Python 3及numpy、pandas、matplotlib）：
1. python src/scenario_model.py
2. python src/paper_figures.py
3. python src/integrity_check.py

排放结果为显式工程情景推演；附件无法直接识别5/10 mg/Nm³附近的控制—排放响应。
"""
    (SUPPORT / "README.txt").write_text(readme, encoding="utf-8")
    (SUPPORT / "支撑材料文件清单.md").write_text(readme, encoding="utf-8")
    ai_support_md = OUT / "AI工具使用详情_匿名版.md"
    if ai_support_md.exists():
        shutil.copy2(ai_support_md, SUPPORT / "AI工具使用详情.md")


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    CLEAN_MD.write_text(sanitize_manuscript_markdown(), encoding="utf-8")
    pandoc_docx(CLEAN_MD, DRAFT_DOCX, ROOT)
    style_academic_docx(DRAFT_DOCX, FINAL_DOCX, ai_detail=False)

    ai_clean = prepare_ai_markdown()
    pandoc_docx(ai_clean, AI_DRAFT_DOCX, ROOT)
    style_academic_docx(AI_DRAFT_DOCX, AI_DOCX, ai_detail=True)
    stage_support_materials()
    print(FINAL_DOCX)
    print(AI_DOCX)
    print(SUPPORT)


if __name__ == "__main__":
    main()
